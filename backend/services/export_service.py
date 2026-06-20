"""
Export service — converts query results into Excel, Image, PDF, and Word files.
"""

import io
import re

import matplotlib
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Font as XLFont, PatternFill
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

matplotlib.use("Agg")

# ── Shared helpers ─────────────────────────────────────────────

_WRAP_AT = 40  # characters before wrapping a cell value


def _truncate(val: str, limit: int = _WRAP_AT) -> str:
    return val if len(val) <= limit else val[:limit - 1] + "…"


def _col_char_widths(df: pd.DataFrame) -> list[int]:
    """Return the max character width per column (header + all values, capped at 40)."""
    widths = []
    for col in df.columns:
        col_vals = df[col].astype(str)
        max_val  = col_vals.str.len().max() if len(df) else 0
        widths.append(min(max(len(str(col)), int(max_val)), _WRAP_AT))
    return widths


# ── CSV ────────────────────────────────────────────────────────

def to_csv(df: pd.DataFrame) -> io.BytesIO:
    """Export DataFrame as CSV bytes."""
    buffer = io.BytesIO()
    df.to_csv(buffer, index=False)
    buffer.seek(0)
    return buffer


# ── JSON ───────────────────────────────────────────────────────

def to_json(df: pd.DataFrame) -> io.BytesIO:
    """Export DataFrame as JSON (records format) bytes."""
    buffer = io.BytesIO()
    df.to_json(buffer, orient='records', date_format='iso')
    buffer.seek(0)
    return buffer


# ── Excel ──────────────────────────────────────────────────────

def to_excel(df: pd.DataFrame) -> io.BytesIO:
    """Export DataFrame to Excel with auto-fitted column widths."""
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Results")
        ws = writer.sheets["Results"]
        for col_cells in ws.columns:
            width = max(len(str(cell.value)) if cell.value else 0 for cell in col_cells)
            ws.column_dimensions[col_cells[0].column_letter].width = width + 4
    buffer.seek(0)
    return buffer


# ── Image ──────────────────────────────────────────────────────

def to_image(df: pd.DataFrame, fmt: str = "png") -> io.BytesIO:
    """Export DataFrame as a table image (PNG or JPG) with content-fitted columns."""
    str_df  = df.astype(str).apply(lambda col: col.map(_truncate))
    n_rows  = len(str_df)
    n_cols  = len(str_df.columns)
    widths  = _col_char_widths(df)

    # Figure size: width proportional to char widths, height proportional to rows
    char_total = sum(widths)
    fig_w = max(8.0, char_total * 0.13)
    fig_h = max(1.5, n_rows * 0.32 + 1.0)

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.axis("off")

    tbl = ax.table(
        cellText=str_df.values,
        colLabels=str_df.columns.tolist(),
        cellLoc="left",
        loc="upper left",
        bbox=[0, 0, 1, 1],
    )
    tbl.auto_set_font_size(False)
    tbl.set_fontsize(9)

    # Set each column width proportionally
    total_w = sum(widths)
    for col_idx, w in enumerate(widths):
        col_frac = w / total_w
        for row_idx in range(n_rows + 1):
            cell = tbl[row_idx, col_idx]
            cell.set_width(col_frac)

    # Styling
    for (row, col), cell in tbl.get_celld().items():
        cell.set_edgecolor("#e2e8f0")
        cell.set_linewidth(0.5)
        cell.PAD = 0.06
        if row == 0:
            cell.set_facecolor("#3b82f6")
            cell.set_text_props(fontweight="bold", color="white", fontsize=9)
        else:
            cell.set_facecolor("#ffffff" if row % 2 == 1 else "#f8fafc")
            cell.set_text_props(color="#1e293b", fontsize=8.5)

    fig.patch.set_facecolor("#ffffff")
    buffer = io.BytesIO()
    fig.savefig(buffer, format=fmt, dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    plt.close(fig)
    buffer.seek(0)
    return buffer


# ── PDF ────────────────────────────────────────────────────────

def to_pdf(df: pd.DataFrame, title: str = "Query Results") -> io.BytesIO:
    """Export DataFrame as a formatted PDF table with content-aware column widths."""
    buffer    = io.BytesIO()
    page_size = landscape(A4) if len(df.columns) > 6 else A4
    usable_w  = page_size[0] - 60          # 30pt margins each side

    char_widths = _col_char_widths(df)
    total_chars = sum(char_widths)
    col_widths  = [usable_w * (w / total_chars) for w in char_widths]

    # Paragraph styles for wrapped cell text
    styles     = getSampleStyleSheet()
    hdr_style  = ParagraphStyle(
        "HdrCell",
        parent=styles["Normal"],
        fontSize=8,
        fontName="Helvetica-Bold",
        textColor=colors.white,
        leading=11,
        wordWrap="LTR",
    )
    cell_style = ParagraphStyle(
        "DataCell",
        parent=styles["Normal"],
        fontSize=7.5,
        fontName="Helvetica",
        textColor=colors.HexColor("#1e293b"),
        leading=10,
        wordWrap="LTR",
    )

    header_row = [Paragraph(str(c), hdr_style) for c in df.columns]
    data_rows  = [
        [Paragraph(_truncate(str(v), 60), cell_style) for v in row]
        for row in df.itertuples(index=False)
    ]
    table_data = [header_row] + data_rows

    tbl = Table(table_data, colWidths=col_widths, repeatRows=1,
                hAlign="LEFT", splitByRow=True)
    tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor("#3b82f6")),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
        ("GRID",          (0, 0), (-1, -1), 0.4, colors.HexColor("#e2e8f0")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN",         (0, 0), (-1, -1), "LEFT"),
        ("TOPPADDING",    (0, 0), (-1, 0),  6),
        ("BOTTOMPADDING", (0, 0), (-1, 0),  6),
        ("TOPPADDING",    (0, 1), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 4),
        ("LEFTPADDING",   (0, 0), (-1, -1), 5),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 5),
    ]))

    title_style = ParagraphStyle(
        "Title", parent=styles["Title"], fontSize=14,
        textColor=colors.HexColor("#1e293b"), spaceAfter=8,
    )
    doc = SimpleDocTemplate(
        buffer, pagesize=page_size,
        leftMargin=30, rightMargin=30, topMargin=30, bottomMargin=30,
    )
    doc.build([Paragraph(title, title_style), Spacer(1, 10), tbl])
    buffer.seek(0)
    return buffer


# ── Word ───────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Apply a background fill colour to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tc_pr.append(shd)


def _set_col_width(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tcW   = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    tc_pr.append(tcW)


def to_word(df: pd.DataFrame, title: str = "Query Results") -> io.BytesIO:
    """Export DataFrame as a formatted Word document with fitted column widths."""
    doc = Document()

    # Page margins — narrow for more table space
    for section in doc.sections:
        section.left_margin   = Twips(720)   # 0.5 in
        section.right_margin  = Twips(720)
        section.top_margin    = Twips(720)
        section.bottom_margin = Twips(720)

    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    char_widths = _col_char_widths(df)
    total_chars = sum(char_widths)
    # Usable page width ≈ 12240 twips (letter) minus margins (1440)
    usable_twips = 12240 - 1440
    col_twips    = [int(usable_twips * (w / total_chars)) for w in char_widths]

    tbl = doc.add_table(rows=1, cols=len(df.columns))
    tbl.style = "Table Grid"

    # Header row
    for i, col in enumerate(df.columns):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, "3b82f6")
        _set_col_width(cell, col_twips[i])
        para = cell.paragraphs[0]
        para.alignment = 1          # WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run(str(col))
        run.bold             = True
        run.font.size        = Pt(9)
        run.font.color.rgb   = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, (_, row) in enumerate(df.iterrows()):
        row_cells = tbl.add_row().cells
        bg        = "ffffff" if r_idx % 2 == 0 else "f8fafc"
        for i, val in enumerate(row):
            cell = row_cells[i]
            _set_cell_bg(cell, bg)
            _set_col_width(cell, col_twips[i])
            para = cell.paragraphs[0]
            para.alignment = 0          # LEFT
            run  = para.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ── Chart export (image + insights + underlying data) ───────────

def _strip_markdown(text: str) -> str:
    """Remove markdown bold markers, returning plain text."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text or "")


def _md_to_reportlab(text: str) -> str:
    """Convert a small markdown subset (bold only) to ReportLab Paragraph markup."""
    text = (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)


def _add_markdown_paragraph(doc: Document, text: str, size: int = 10) -> None:
    """Add a bulleted paragraph to a docx document, rendering **bold** spans as bold runs."""
    para = doc.add_paragraph(style="List Bullet")
    for part in re.split(r"(\*\*.+?\*\*)", text or ""):
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        run = para.add_run(part[2:-2] if is_bold else part)
        run.bold = is_bold
        run.font.size = Pt(size)


def _fmt_chart_value(v) -> str:
    try:
        fv = float(v)
        return f"{int(fv):,}" if fv == int(fv) else f"{fv:,.2f}"
    except (TypeError, ValueError):
        return str(v)


def to_chart_pdf(title, image_bytes, x_label, y_label, labels, data, series_label, insights) -> io.BytesIO:
    """Export a chart (rendered image + insights + underlying data) as a PDF report."""
    buffer = io.BytesIO()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ChartTitle", parent=styles["Title"], fontSize=16,
        textColor=colors.HexColor("#1e293b"), spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "ChartHeading", parent=styles["Heading2"], fontSize=12,
        textColor=colors.HexColor("#1e293b"), spaceBefore=14, spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ChartBody", parent=styles["Normal"], fontSize=10,
        textColor=colors.HexColor("#334155"), leading=14,
    )

    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40,
    )
    story = [Paragraph(title, title_style)]

    img_w, img_h = PILImage.open(io.BytesIO(image_bytes)).size
    max_w = A4[0] - 80
    display_w = min(max_w, img_w)
    display_h = display_w * img_h / img_w
    story.append(RLImage(io.BytesIO(image_bytes), width=display_w, height=display_h))

    if insights:
        story.append(Paragraph("Key Insights", heading_style))
        for insight in insights:
            story.append(Paragraph(f"&bull; {_md_to_reportlab(insight)}", body_style))

    if labels and data:
        story.append(Paragraph("Chart Data", heading_style))
        table_data = [[x_label or "Category", series_label or y_label or "Value"]]
        table_data += [[str(l), _fmt_chart_value(v)] for l, v in zip(labels, data)]
        tbl = Table(table_data, hAlign="LEFT")
        tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor("#3b82f6")),
            ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
            ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
            ("GRID",          (0, 0), (-1, -1), 0.4, colors.HexColor("#e2e8f0")),
            ("FONTSIZE",      (0, 0), (-1, -1), 9),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
        ]))
        story.append(tbl)

    doc.build(story)
    buffer.seek(0)
    return buffer


def to_chart_word(title, image_bytes, x_label, y_label, labels, data, series_label, insights) -> io.BytesIO:
    """Export a chart (rendered image + insights + underlying data) as a Word document."""
    doc = Document()
    for section in doc.sections:
        section.left_margin   = Twips(720)
        section.right_margin  = Twips(720)
        section.top_margin    = Twips(720)
        section.bottom_margin = Twips(720)

    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    doc.add_picture(io.BytesIO(image_bytes), width=Inches(6.0))

    if insights:
        doc.add_heading("Key Insights", level=2)
        for insight in insights:
            _add_markdown_paragraph(doc, insight)

    if labels and data:
        doc.add_heading("Chart Data", level=2)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        for i, col in enumerate([x_label or "Category", series_label or y_label or "Value"]):
            cell = tbl.rows[0].cells[i]
            _set_cell_bg(cell, "3b82f6")
            run = cell.paragraphs[0].add_run(str(col))
            run.bold           = True
            run.font.size      = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for r_idx, (label, value) in enumerate(zip(labels, data)):
            row_cells = tbl.add_row().cells
            bg = "ffffff" if r_idx % 2 == 0 else "f8fafc"
            for i, val in enumerate((label, _fmt_chart_value(value))):
                cell = row_cells[i]
                _set_cell_bg(cell, bg)
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size      = Pt(9)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def to_chart_excel(title, image_bytes, x_label, y_label, labels, data, series_label, insights) -> io.BytesIO:
    """Export a chart (embedded image + underlying data + insights) as an Excel workbook."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Chart"
    ws["A1"] = title
    ws["A1"].font = XLFont(size=14, bold=True, color="1E293B")

    xl_img = XLImage(io.BytesIO(image_bytes))
    xl_img.width, xl_img.height = 640, 380
    ws.add_image(xl_img, "A3")

    ws_data = wb.create_sheet("Chart Data")
    if not labels and data and isinstance(data[0], dict) and "x" in data[0] and "y" in data[0]:
        # Scatter points are {x, y} pairs rather than label/value rows
        header = [x_label or "X", y_label or "Y"]
        rows = [[p.get("x"), p.get("y")] for p in data]
    else:
        header = [x_label or "Category", series_label or y_label or "Value"]
        rows = [list(pair) for pair in zip(labels, data)]

    ws_data.append(header)
    for cell in ws_data[1]:
        cell.font = XLFont(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="3B82F6")
    for row in rows:
        ws_data.append(row)
    for col_cells in ws_data.columns:
        width = max((len(str(c.value)) if c.value is not None else 0) for c in col_cells)
        ws_data.column_dimensions[col_cells[0].column_letter].width = width + 4

    if insights:
        ws_insights = wb.create_sheet("Insights")
        ws_insights["A1"] = "Key Insights"
        ws_insights["A1"].font = XLFont(size=12, bold=True, color="1E293B")
        for i, insight in enumerate(insights, start=2):
            ws_insights.cell(row=i, column=1, value=_strip_markdown(insight))
        ws_insights.column_dimensions["A"].width = 100

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer
